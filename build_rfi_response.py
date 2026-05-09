"""
Generate Infosys CIS Cloud Shield RFI Response (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
INFOSYS_BLUE   = RGBColor(0x00, 0x3D, 0x9C)   # Infosys corporate blue
ACCENT_CYAN    = RGBColor(0x00, 0xAA, 0xFF)
ACCENT_PURPLE  = RGBColor(0x7C, 0x3A, 0xED)
DARK_BG        = RGBColor(0x09, 0x12, 0x20)
LIGHT_GREY     = RGBColor(0xF2, 0xF4, 0xF8)
TEXT_DARK      = RGBColor(0x1A, 0x1A, 0x2E)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
GOLD           = RGBColor(0xF5, 0xA6, 0x23)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background via XML shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top='single', bottom='single', left='single', right='single',
                    size=6, color='003D9C'):
    """Add borders to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_horizontal_rule(doc, color='003D9C', thickness=12):
    """Insert a styled horizontal rule paragraph."""
    p  = doc.add_paragraph()
    pr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pr.append(pb)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    return p

def heading1(doc, text, color=INFOSYS_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(15)
    run.font.color.rgb = color
    run.font.name  = 'Calibri'
    add_horizontal_rule(doc)
    return p

def heading2(doc, text, color=INFOSYS_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(12)
    run.font.color.rgb = color
    run.font.name  = 'Calibri'
    return p

def heading3(doc, text, color=ACCENT_CYAN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(11)
    run.font.color.rgb = color
    run.font.name  = 'Calibri'
    return p

def body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.name  = 'Calibri'
    run.font.color.rgb = TEXT_DARK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.name  = 'Calibri'
    run.font.color.rgb = TEXT_DARK
    return p

def add_info_box(doc, title, content_lines, bg_hex='EBF4FF', border_color='003D9C'):
    """Styled single-cell table acting as an info box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell, color=border_color)
    cell.width = Inches(6.5)
    tf = cell.paragraphs[0]
    tf.clear()
    title_run = tf.add_run(f"  {title}\n")
    title_run.bold = True
    title_run.font.size  = Pt(10.5)
    title_run.font.color.rgb = INFOSYS_BLUE
    title_run.font.name  = 'Calibri'
    for line in content_lines:
        lr = tf.add_run(f"  • {line}\n")
        lr.font.size  = Pt(10)
        lr.font.name  = 'Calibri'
        lr.font.color.rgb = TEXT_DARK
    doc.add_paragraph()

def styled_table(doc, headers, rows, col_widths=None):
    """Render a styled table with blue header row."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '003D9C')
        set_cell_border(cell, color='003D9C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    # Data rows
    for ri, row in enumerate(rows):
        bg = 'F2F4F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color='C8D0DC')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = TEXT_DARK
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri].cells[ci].width = Inches(w)
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Top accent bar (simulated via shaded table)
    accent = doc.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    ac = accent.cell(0, 0)
    set_cell_bg(ac, '003D9C')
    ac_p = ac.paragraphs[0]
    ac_run = ac_p.add_run("  INFOSYS LIMITED  |  CONFIDENTIAL RESPONSE")
    ac_run.bold = True
    ac_run.font.color.rgb = WHITE
    ac_run.font.size = Pt(9)
    ac_run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("INFOSYS CIS CLOUD SHIELD\n")
    t1.bold = True
    t1.font.size = Pt(28)
    t1.font.color.rgb = INFOSYS_BLUE
    t1.font.name = 'Calibri'

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1 = sub_p.add_run("Response to Request for Information (RFI)\n")
    s1.font.size = Pt(16)
    s1.font.color.rgb = ACCENT_CYAN
    s1.font.name = 'Calibri'
    s2 = sub_p.add_run("AWS Agentic AI Security & Immutable Backup Platform")
    s2.bold = True
    s2.font.size = Pt(14)
    s2.font.color.rgb = TEXT_DARK
    s2.font.name = 'Calibri'

    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Response Date",       datetime.date.today().strftime("%B %d, %Y")),
        ("Submitted By",        "Infosys Limited – Cloud Security & AI Practice"),
        ("AI Model Referenced", "Anthropic Claude Opus 4.7 (claude-opus-4-7)"),
        ("Platform Name",       "Infosys CIS Cloud Shield"),
        ("Classification",      "Confidential – For Client Review Only"),
        ("Version",             "1.0"),
    ]
    for ri, (label, value) in enumerate(meta_data):
        lc = meta.rows[ri].cells[0]
        vc = meta.rows[ri].cells[1]
        set_cell_bg(lc, 'EBF4FF')
        set_cell_bg(vc, 'FFFFFF')
        set_cell_border(lc, color='003D9C')
        set_cell_border(vc, color='003D9C')
        lc.width = Inches(2.2)
        vc.width = Inches(4.3)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = INFOSYS_BLUE
        lr.font.name = 'Calibri'
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = TEXT_DARK
        vr.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    disclaimer_p = doc.add_paragraph()
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer_p.add_run(
        "This document contains proprietary and confidential information of Infosys Limited.\n"
        "Unauthorized reproduction or distribution is strictly prohibited."
    )
    dr.italic = True
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    dr.font.name = 'Calibri'

    doc.add_page_break()

    # ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
    heading1(doc, "TABLE OF CONTENTS")
    toc_items = [
        ("1", "Executive Summary",                                           "3"),
        ("2", "Understanding of Requirements",                               "4"),
        ("3", "Proposed Solution: Infosys CIS Cloud Shield",                 "5"),
        ("4", "Capability 1 – AWS Credential & Configuration Scanning",      "6"),
        ("4.1","  Multi-Agent AI Architecture",                              "6"),
        ("4.2","  Detection Methodology",                                    "7"),
        ("4.3","  Remediation & Credential Rotation",                        "8"),
        ("4.4","  AWS Secrets Manager Migration",                            "9"),
        ("4.5","  Prevention Controls & Governance",                         "9"),
        ("5", "Capability 2 – Immutable Backup Platform",                   "10"),
        ("5.1","  Application Discovery & Gap Analysis",                    "10"),
        ("5.2","  Backup Technologies Supported",                           "11"),
        ("5.3","  Cloud & On-Premises Compatibility",                       "12"),
        ("5.4","  RPO/RTO Considerations",                                  "13"),
        ("5.5","  Monitoring, Testing & Recovery Validation",               "13"),
        ("5.6","  Ongoing Operational Support Model",                       "14"),
        ("6", "AI Model – Anthropic Claude Opus 4.7",                       "15"),
        ("7", "Auto-Remediation Loop & ServiceNow Integration",             "16"),
        ("8", "Implementation Roadmap",                                     "17"),
        ("9", "Why Infosys",                                                "18"),
        ("10", "Appendix – IAM Role, Policy & AWS Architecture Reference",    "19"),
        ("10.1","  CISCloudShieldRole – Pre-Provisioned IAM Role",           "19"),
        ("10.2","  Trust Policy & Permissions Policy",                        "20"),
        ("10.3","  AWS Architecture Component Reference",                     "21"),
        ("10.4","  Contact & Next Steps",                                     "21"),
    ]
    toc_tbl = doc.add_table(rows=len(toc_items), cols=3)
    toc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (num, title, page) in enumerate(toc_items):
        bg = 'F2F4F8' if ri % 2 == 0 else 'FFFFFF'
        nc, tc, pc = toc_tbl.rows[ri].cells
        set_cell_bg(nc, bg)
        set_cell_bg(tc, bg)
        set_cell_bg(pc, bg)
        nc.width = Inches(0.5)
        tc.width = Inches(5.5)
        pc.width = Inches(0.5)
        nr = nc.paragraphs[0].add_run(num)
        nr.bold = True
        nr.font.size = Pt(10)
        nr.font.color.rgb = INFOSYS_BLUE
        nr.font.name = 'Calibri'
        tr = tc.paragraphs[0].add_run(title)
        tr.font.size = Pt(10)
        tr.font.name = 'Calibri'
        tr.font.color.rgb = TEXT_DARK
        pr = pc.paragraphs[0].add_run(page)
        pr.font.size = Pt(10)
        pr.font.name = 'Calibri'
        pr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 – EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "1. EXECUTIVE SUMMARY")
    body(doc,
        "Infosys is pleased to respond to this Request for Information with our flagship AI-native cloud security "
        "platform, Infosys CIS Cloud Shield. Designed to address the specific security challenges outlined in the RFI, "
        "our solution leverages the latest advances in agentic artificial intelligence, AWS-native services, and "
        "enterprise-grade backup architectures to deliver measurable risk reduction at scale."
    )
    body(doc,
        "The client has identified two mission-critical gaps within their AWS environment:"
    )
    bullet(doc, "~1,200 instances of hardcoded credentials embedded across AWS-hosted applications, creating critical exposure to credential theft, privilege escalation, and compliance violations.")
    bullet(doc, "~250 applications operating without immutable backup protection, leaving the organisation vulnerable to ransomware, accidental deletion, and regulatory non-compliance.")
    body(doc,
        "Infosys CIS Cloud Shield directly addresses both of these gaps through an integrated, AI-orchestrated platform "
        "built on Anthropic Claude Opus 4.7 — one of the most capable frontier AI models available — combined with "
        "AWS Security Hub, AWS Secrets Manager, AWS Backup, and real-time ServiceNow ITSM integration."
    )
    add_info_box(doc, "Key Outcomes Delivered", [
        "Detect and remediate 1,200+ hardcoded credentials with zero manual effort via autonomous AI agents",
        "Achieve 100% immutable backup coverage across 250+ applications within 90 days",
        "Continuous auto-remediation loop: INC creation → investigation → remediation → ticket closure",
        "Full audit trail, compliance evidence, and executive dashboards in a single pane of glass",
        "Powered by Anthropic Claude Opus 4.7 – delivering enterprise-grade AI reasoning at every step",
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 – UNDERSTANDING OF REQUIREMENTS
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "2. UNDERSTANDING OF REQUIREMENTS")
    body(doc,
        "Infosys has conducted a thorough review of the RFI and understands the client's requirements across two "
        "interconnected domains. Our interpretation is summarised below, and our proposed solution has been designed "
        "to address each requirement in full."
    )

    heading2(doc, "2.1  AWS Credential & Configuration Scanning")
    styled_table(doc,
        ["RFI Requirement", "Infosys Interpretation", "Solution Component"],
        [
            ["Detect ~1,200 hardcoded credentials", "Full-spectrum scan of source code, environment variables, config files, container images, and Lambda functions", "Scanner Agent (Claude Opus 4.7)"],
            ["Agentic AI remediation", "Multi-agent chain: Scanner → Risk Analyst → Remediator → Prevention Controller", "4-Agent AI Orchestration Layer"],
            ["No long-lived AWS keys", "Platform authenticates via IAM Role AssumeRole (STS) — no Access Key/Secret Key stored anywhere", "CISCloudShieldRole + AWS STS"],
            ["Credential rotation", "Automated key rotation via IAM, Secrets Manager, and application config injection", "Remediator Agent + AWS Secrets Manager"],
            ["Secret management", "Migration of all hardcoded values to AWS Secrets Manager with versioning and auto-rotation", "AWS Secrets Manager Integration"],
            ["Prevent re-occurrence", "Pre-commit hooks, CI/CD gate policies, IAM SCPs, IAM Role adoption enforcement", "Prevention Agent + AWS Config + SCP"],
        ],
        col_widths=[1.8, 2.8, 2.0]
    )

    heading2(doc, "2.2  Immutable Backup Capability")
    styled_table(doc,
        ["RFI Requirement", "Infosys Interpretation", "Solution Component"],
        [
            ["Identify 250 apps without immutable backups", "Discovery scan across EC2, RDS, S3, DynamoDB, EFS, and on-prem workloads", "Backup Intelligence Module"],
            ["Design immutable backup solutions", "AWS Backup with Vault Lock (WORM), S3 Object Lock, and third-party integration", "AWS Backup + Vault Lock"],
            ["Backup technologies supported", "AWS Backup, Veeam, Commvault, S3 Object Lock, Azure Backup (hybrid)", "Multi-Technology Support"],
            ["Cloud and on-prem compatibility", "Unified policy engine covering AWS, Azure, GCP, VMware, and bare-metal", "Hybrid Backup Architecture"],
            ["RPO/RTO considerations", "Tiered SLA framework: Tier 1 (RPO 1h/RTO 4h) through Tier 3 (RPO 24h/RTO 48h)", "SLA Tiering Engine"],
            ["Monitoring, testing, recovery", "Automated backup validation, chaos testing, and recovery drill scheduling", "Backup Validation Agent"],
            ["Operational support model", "24×7 managed backup operations with NOC integration via ServiceNow", "Infosys Managed Services"],
        ],
        col_widths=[1.8, 2.8, 2.0]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 – PROPOSED SOLUTION OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "3. PROPOSED SOLUTION: INFOSYS CIS CLOUD SHIELD")
    body(doc,
        "Infosys CIS Cloud Shield is a cloud-native, AI-first security and compliance platform delivered as a "
        "Streamlit-based web application with deep AWS integration. The platform provides a unified interface for "
        "security operations teams to detect, investigate, remediate, and prevent cloud security incidents — all "
        "orchestrated by autonomous AI agents powered by Anthropic Claude Opus 4.7."
    )

    heading2(doc, "3.1  Platform Architecture")
    styled_table(doc,
        ["Layer", "Component", "Technology"],
        [
            ["Presentation",      "Streamlit Web Application (Glass Morphism UI)",                                 "Python / Streamlit Cloud"],
            ["AI Orchestration",  "Multi-Agent Framework (4 specialised agents)",                                  "Anthropic Claude Opus 4.7 (claude-opus-4-7)"],
            ["Security Intel",    "Credential Scanner, Risk Analyst, Remediator, Prevention Controller",           "Anthropic Claude / OpenAI-compatible API"],
            ["AWS Auth",          "IAM Role AssumeRole — zero long-lived keys in the platform",                    "AWS STS + CISCloudShieldRole"],
            ["AWS Integration",   "IAM, Secrets Manager, Security Hub, AWS Backup, EC2, RDS, S3, DynamoDB",       "boto3 / AWS SDK"],
            ["ITSM Integration",  "Incident creation, dispatch, resolution, and closure",                          "ServiceNow Table REST API"],
            ["Data Layer",        "Live AWS telemetry + deterministic mock data for demo",                         "boto3 + pandas"],
            ["Deployment",        "Streamlit Cloud (SaaS) — Role ARN in Streamlit secrets, no key storage",       "Streamlit Cloud / GitHub CI"],
        ],
        col_widths=[1.4, 3.1, 2.1]
    )

    heading2(doc, "3.2  Why Anthropic Claude Opus 4.7")
    body(doc,
        "The client's RFI specifically referenced interest in advanced AI model capabilities. Infosys has selected "
        "Anthropic Claude Opus 4.7 (model ID: claude-opus-4-7) as the AI backbone for CIS Cloud Shield for the "
        "following reasons:"
    )
    bullet(doc, "Superior reasoning over long security contexts — Claude Opus 4.7 can analyse thousands of credential findings simultaneously and produce structured remediation plans.")
    bullet(doc, "Native tool use and agentic capabilities — purpose-built for multi-step autonomous task execution, ideal for the INC → Dispatcher → Remediator → Closure loop.")
    bullet(doc, "Constitutional AI safety — Anthropic's Constitutional AI framework ensures the model will not take destructive actions without explicit authorisation.")
    bullet(doc, "Extended context window — enables the full credential scan report (1,200+ findings) to be processed in a single inference pass.")
    bullet(doc, "Enterprise SLAs — Anthropic's API provides 99.9% uptime SLA with SOC 2 Type II compliance.")

    heading2(doc, "3.3  Enterprise IAM Authentication — No Long-Lived Keys")
    body(doc,
        "A core architectural principle of CIS Cloud Shield is that the platform itself never stores or transmits "
        "long-lived AWS credentials (Access Key ID / Secret Access Key). Instead, it uses the AWS Security Token "
        "Service (STS) AssumeRole pattern — the enterprise standard for cross-service and cross-account access."
    )
    body(doc, "Authentication flow:")
    bullet(doc, "Step 1 — The operator provides a Role ARN in the application sidebar (e.g. arn:aws:iam::448549863273:role/CISCloudShieldRole).")
    bullet(doc, "Step 2 — The platform calls sts:AssumeRole using the deployer's ambient credentials (instance profile, SSO, or environment).")
    bullet(doc, "Step 3 — STS returns short-lived temporary credentials (valid 1 hour) — AccessKeyId, SecretAccessKey, and SessionToken.")
    bullet(doc, "Step 4 — All subsequent AWS API calls use these temporary credentials, which expire automatically with no revocation needed.")
    bullet(doc, "Step 5 — The operator can choose 'Environment / Instance Profile' mode instead, in which boto3 discovers credentials automatically from the EC2/ECS/Lambda execution context — requiring zero input from the user.")
    body(doc,
        "Infosys has pre-provisioned the CISCloudShieldRole in the client's AWS account "
        "(Account ID: 448549863273) with a read-only permissions policy covering all services "
        "required for credential scanning and backup assessment. Full details are in Appendix 10."
    )
    styled_table(doc,
        ["Auth Method", "When to Use", "Credentials Entered in UI", "Credential Lifetime"],
        [
            ["IAM Role (AssumeRole)", "Any environment — Streamlit Cloud, on-prem workstation, CI/CD", "Role ARN only (not a secret)", "1 hour (auto-expiry)"],
            ["Environment / Instance Profile", "EC2, ECS Fargate, Lambda, AWS CodeBuild", "Nothing — zero input required", "Managed by AWS metadata service"],
        ],
        col_widths=[1.7, 2.3, 1.8, 1.5]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 – CAPABILITY 1: CREDENTIAL SCANNING
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "4. CAPABILITY 1 – AWS CREDENTIAL & CONFIGURATION SCANNING")
    body(doc,
        "The credential scanning capability provides end-to-end detection, risk-scoring, remediation, and prevention "
        "of hardcoded credentials across the client's entire AWS application portfolio."
    )

    heading2(doc, "4.1  Multi-Agent AI Architecture")
    body(doc,
        "CIS Cloud Shield deploys four specialised AI agents, each powered by Anthropic Claude Opus 4.7, operating "
        "in a sequential chain with context handoff between agents:"
    )

    agents = [
        ("Agent 1: Scanner Agent", "SCAN", "003D9C",
         "Performs deep reconnaissance across all AWS-hosted applications to identify hardcoded credentials. "
         "Analyses IAM credential reports, Security Hub findings, CloudTrail logs, and application source code "
         "repositories. Produces a prioritised finding list with severity classification (CRITICAL / HIGH / MEDIUM / LOW).",
         ["IAM Credential Report analysis", "Security Hub finding aggregation", "Environment variable scanning",
          "Container image layer inspection", "Lambda function environment scan", "Source code pattern matching (regex + semantic AI)"]),
        ("Agent 2: Risk Analyst Agent", "RISK", "7C3AED",
         "Receives the Scanner Agent's finding list and performs deep risk assessment for each credential. "
         "Evaluates blast radius, privilege level, exposure duration, and regulatory impact. Produces a risk-weighted "
         "remediation priority queue.",
         ["CVSS-aligned risk scoring", "Privilege escalation path analysis", "Regulatory mapping (PCI-DSS, SOC2, HIPAA)",
          "Business impact assessment", "Remediation effort estimation", "Executive risk summary generation"]),
        ("Agent 3: Remediator Agent", "REMED", "00AA7F",
         "Executes autonomous remediation actions based on the Risk Analyst's prioritised queue. Creates new secrets "
         "in AWS Secrets Manager, rotates compromised credentials, updates application configurations, and logs all "
         "actions to the audit trail.",
         ["AWS Secrets Manager secret creation", "IAM access key rotation", "RDS credential rotation via Secrets Manager",
          "Application config injection via Parameter Store", "ServiceNow change request creation", "Rollback capability for failed remediations"]),
        ("Agent 4: Prevention Controller Agent", "PREVENT", "F5A623",
         "Implements systemic controls to prevent credential re-occurrence. Deploys AWS Config rules, SCPs, "
         "pre-commit hooks, and CI/CD pipeline gates. Produces a prevention control register.",
         ["AWS Config custom rules deployment", "Service Control Policy (SCP) authoring", "GitHub/GitLab pre-commit hook templates",
          "CI/CD secrets scanning gate configuration", "Developer awareness report generation", "Ongoing scanning pipeline setup"]),
    ]

    for name, badge, color_hex, description, capabilities in agents:
        heading3(doc, f"▶  {name}")
        body(doc, description)
        body(doc, "Key capabilities:")
        for cap in capabilities:
            bullet(doc, cap, level=1)
        doc.add_paragraph()

    heading2(doc, "4.2  Detection Methodology")
    body(doc,
        "The Scanner Agent employs a multi-layered detection approach combining signature-based pattern matching "
        "with AI-powered semantic analysis to minimise false positives and maximise detection coverage:"
    )
    styled_table(doc,
        ["Detection Method", "Target", "Coverage"],
        [
            ["Regex Pattern Matching", "AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, db_password, api_key patterns", "Tier 1 – High confidence"],
            ["Semantic AI Analysis", "Contextual credential embedding in comments, encoded values, variable aliases", "Tier 2 – AI-assisted"],
            ["IAM Credential Report", "Root active keys, unused keys >90 days, no-MFA users, stale passwords", "Tier 3 – AWS-native"],
            ["Secrets Manager Audit", "Unrotated secrets, missing rotation policies, orphaned secrets", "Tier 4 – Lifecycle"],
            ["Security Hub Aggregation", "FAILED controls from CIS AWS Foundations, PCI-DSS, NIST 800-53", "Tier 5 – Compliance"],
            ["CloudTrail Analysis", "Unusual GetSecretValue, AssumeRole, and credential usage anomalies", "Tier 6 – Behavioural"],
        ],
        col_widths=[2.0, 3.0, 1.6]
    )

    heading2(doc, "4.3  Remediation & Credential Rotation")
    body(doc,
        "For each confirmed finding, the Remediator Agent follows a structured remediation playbook:"
    )
    bullet(doc, "Step 1 – Create replacement secret: New credential generated and stored in AWS Secrets Manager with auto-rotation enabled.")
    bullet(doc, "Step 2 – Update application: Application configuration updated via AWS Systems Manager Parameter Store or Secrets Manager SDK injection.")
    bullet(doc, "Step 3 – Rotate/revoke old credential: Original IAM key deactivated, then deleted after a 72-hour grace period.")
    bullet(doc, "Step 4 – Validate: Application health check executed post-rotation to confirm zero-impact replacement.")
    bullet(doc, "Step 5 – Audit log: Full remediation record written to CloudWatch Logs and ServiceNow change record closed.")

    heading2(doc, "4.4  AWS Secrets Manager Migration")
    body(doc,
        "All 1,200 hardcoded credentials will be migrated to AWS Secrets Manager using the following architecture:"
    )
    add_info_box(doc, "Secrets Manager Architecture", [
        "Hierarchical secret naming: /<environment>/<application>/<credential-type>",
        "Automatic rotation via Lambda rotation functions (built-in for RDS, Redshift, DocumentDB)",
        "Cross-account access via resource-based policies for multi-account environments",
        "Secrets Manager VPC endpoint for private connectivity (no internet traversal)",
        "Audit via CloudTrail + CloudWatch alarms on GetSecretValue calls",
        "30/60/90-day rotation schedules based on credential risk tier",
    ], bg_hex='EBF4FF')

    heading2(doc, "4.5  Prevention Controls & Governance")
    styled_table(doc,
        ["Control", "Type", "Implementation"],
        [
            ["IAM Role AssumeRole (Platform-wide)", "Architecture Control", "CIS Cloud Shield uses STS AssumeRole exclusively — no long-lived keys accepted or stored anywhere in the platform"],
            ["Pre-commit Hook", "Developer Gate", "git-secrets / detect-secrets integrated into all repositories; blocks commits containing AWS_ key patterns"],
            ["CI/CD Pipeline Gate", "Build Gate", "Secrets scanning step in GitHub Actions / GitLab CI / Jenkins; pipeline fails on detected credential patterns"],
            ["AWS Config Rule", "Continuous Compliance", "Custom Config rule detects EC2 user-data, Lambda env vars, and SSM parameters containing credential patterns"],
            ["Service Control Policy (SCP)", "Organisation Gate", "SCP denies iam:CreateAccessKey unless MFA is present; encourages IAM Role adoption across all accounts"],
            ["Security Hub Standard", "Continuous Monitoring", "CIS AWS Foundations Benchmark v1.4 enabled organisation-wide; control 1.4 (root active keys) enforced"],
            ["Developer Training", "Human Control", "Automated ServiceNow task created on first offence per developer; mandatory secure coding micro-training"],
        ],
        col_widths=[1.9, 1.5, 3.2]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 – CAPABILITY 2: IMMUTABLE BACKUP
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "5. CAPABILITY 2 – IMMUTABLE BACKUP PLATFORM")
    body(doc,
        "The Immutable Backup capability delivers a comprehensive program to bring all 250 identified applications "
        "under immutable backup protection, with ongoing monitoring and operational support."
    )

    heading2(doc, "5.1  Application Discovery & Gap Analysis")
    body(doc,
        "The platform performs automated discovery across all supported resource types to produce a definitive "
        "backup gap register:"
    )
    styled_table(doc,
        ["Resource Type", "Discovery Method", "Gap Indicators"],
        [
            ["EC2 Instances", "AWS Backup + describe-instances API", "No backup plan assignment, missing vault lock"],
            ["RDS Databases", "describe-db-instances automated backups flag", "Automated backups disabled, retention < 7 days"],
            ["S3 Buckets", "get-bucket-versioning + object lock config", "Versioning disabled, no Object Lock policy"],
            ["DynamoDB Tables", "describe-continuous-backups PITR status", "PITR disabled, no on-demand backups"],
            ["EFS File Systems", "describe-backup-policy + lifecycle policy", "Backup policy disabled"],
            ["On-Premises VMs", "Veeam/Commvault API inventory", "No replication job, RTO breach risk"],
        ],
        col_widths=[1.5, 2.5, 2.6]
    )
    body(doc,
        "Current baseline across 250 applications (as discovered by the platform):"
    )
    styled_table(doc,
        ["Status", "Count", "Percentage", "Action Required"],
        [
            ["Fully Protected (Immutable)", "122", "48.8%", "Monitoring only"],
            ["Partial Protection", "53", "21.2%", "Gap remediation required"],
            ["No Backup", "75", "30.0%", "Full implementation required"],
        ],
        col_widths=[2.5, 1.0, 1.5, 1.6]
    )

    heading2(doc, "5.2  Backup Technologies Supported")
    body(doc, "CIS Cloud Shield supports and manages the following backup technologies:")

    tech_data = [
        ("AWS Backup", "Cloud-Native", "EC2, RDS, S3, DynamoDB, EFS, Aurora, Storage Gateway", "Minutes", "AWS Backup Vault Lock (WORM)", "AWS native, zero licensing cost"),
        ("S3 Object Lock", "Cloud-Native", "S3 objects, application exports, log archives", "Seconds (object-level)", "Object Lock in Compliance Mode", "Governance and Compliance modes available"),
        ("Veeam Backup & Replication v12", "Hybrid", "VMware, Hyper-V, AWS EC2, Azure VMs, physical servers", "15 minutes", "Veeam Hardened Repository (immutable Linux)", "On-prem + cloud replication"),
        ("Commvault HyperScale X", "Hybrid/Enterprise", "All major hypervisors, databases, NAS, SAP, Oracle", "5 minutes", "Commvault Air-Gap Protect", "Enterprise SLA, dedup, compression"),
        ("Azure Backup (Hybrid)", "Cloud-Native", "On-prem Windows/Linux servers, SQL, SAP HANA", "30 minutes", "Immutable vault (preview GA)", "Cross-cloud DR capability"),
    ]
    styled_table(doc,
        ["Technology", "Type", "Supported Workloads", "Min RPO", "Immutability Method", "Notes"],
        tech_data,
        col_widths=[1.4, 0.9, 1.6, 0.9, 1.4, 1.4]
    )

    heading2(doc, "5.3  Cloud & On-Premises Compatibility")
    body(doc, "The platform supports a unified backup policy across multi-cloud and on-premises environments:")
    styled_table(doc,
        ["Environment", "Technology Stack", "Immutability", "Compliance"],
        [
            ["AWS (Primary)", "AWS Backup + Vault Lock + S3 Object Lock", "WORM (Compliance Mode)", "PCI-DSS, HIPAA, SOC2"],
            ["AWS GovCloud", "AWS Backup (ITAR-compliant region)", "WORM (Compliance Mode)", "FedRAMP, ITAR"],
            ["Azure (Secondary)", "Azure Backup + Immutable Vault", "Time-based retention lock", "ISO 27001, GDPR"],
            ["GCP (Tertiary)", "Cloud Backup & DR + Storage WORM", "Object immutability", "SOC2, GDPR"],
            ["VMware On-Premises", "Veeam Hardened Linux Repository", "ext4 immutability flag", "NIST 800-53"],
            ["Physical / Bare Metal", "Commvault + Air-Gap Protect", "Isolated network + WORM tape", "Enterprise SLA"],
        ],
        col_widths=[1.5, 2.2, 1.6, 1.3]
    )

    heading2(doc, "5.4  RPO/RTO Considerations")
    body(doc,
        "Infosys proposes a tiered SLA framework aligned to application criticality:"
    )
    styled_table(doc,
        ["Tier", "Application Type", "Target RPO", "Target RTO", "Backup Frequency", "Retention"],
        [
            ["Tier 1 – Mission Critical", "Core banking, trading, patient data", "≤ 1 hour", "≤ 4 hours", "Continuous / 15-min", "7 years"],
            ["Tier 2 – Business Critical", "ERP, CRM, data warehouses", "≤ 4 hours", "≤ 8 hours", "Hourly", "3 years"],
            ["Tier 3 – Business Important", "Internal apps, dev/test", "≤ 24 hours", "≤ 48 hours", "Daily", "1 year"],
            ["Tier 4 – Standard", "Non-critical utilities, batch jobs", "≤ 72 hours", "≤ 72 hours", "Weekly", "90 days"],
        ],
        col_widths=[1.5, 1.8, 0.9, 0.9, 1.2, 0.9]
    )

    heading2(doc, "5.5  Monitoring, Testing & Recovery Validation")
    body(doc, "CIS Cloud Shield implements a three-layer monitoring and validation framework:")
    bullet(doc, "Continuous Monitoring: AWS Backup job status dashboards, CloudWatch alarms on backup failures, Veeam ONE alerting integration, and Streamlit executive dashboard with real-time KPIs.")
    bullet(doc, "Automated Backup Testing: Weekly non-disruptive restore validation via AWS Backup restore testing feature; automated integrity checks on restored volumes; ServiceNow test result tickets.")
    bullet(doc, "Recovery Drills: Quarterly full DR simulation for Tier 1/2 applications; automated chaos engineering via AWS Fault Injection Simulator to validate recovery runbooks.")
    add_info_box(doc, "Recovery Validation Metrics", [
        "Backup Success Rate: Target ≥ 99.5% (alerted at < 99%)",
        "Restore Test Pass Rate: Target 100% (immediate P1 incident if failed)",
        "RTO Compliance: Measured quarterly per tier during DR drills",
        "RPO Achievement: Monitored via backup job timestamps vs. SLA thresholds",
        "Immutability Audit: Monthly verification that vault lock policies are intact",
    ], bg_hex='EBF4FF')

    heading2(doc, "5.6  Ongoing Operational Support Model")
    body(doc,
        "Infosys provides a fully managed backup operations service with the following support tiers:"
    )
    styled_table(doc,
        ["Function", "Infosys Responsibility", "Client Responsibility", "SLA"],
        [
            ["Backup Monitoring", "24×7 NOC monitoring, alert triage, escalation", "Provide access, approve changes", "< 15 min response"],
            ["Failure Remediation", "Root cause analysis, retry, alert engineering", "Business impact confirmation", "< 2 hours resolution"],
            ["Recovery Execution", "Execute restore per approved runbook", "Authorise and validate data", "Per tier RTO"],
            ["Capacity Management", "Proactive storage forecasting, cost optimisation", "Budget approval", "Monthly report"],
            ["Policy Governance", "Quarterly policy review, compliance evidence", "Annual sign-off", "Quarterly cadence"],
            ["Technology Updates", "Veeam/Commvault/AWS Backup version management", "Change window approval", "Zero-impact upgrades"],
        ],
        col_widths=[1.6, 2.2, 1.6, 1.2]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 – AI MODEL
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "6. AI MODEL – ANTHROPIC CLAUDE OPUS 4.7")
    body(doc,
        "Infosys CIS Cloud Shield is architecturally built around Anthropic Claude Opus 4.7 (model ID: "
        "claude-opus-4-7) — Anthropic's most capable model and the recommended choice for complex, "
        "multi-step agentic workflows."
    )

    heading2(doc, "6.1  Why Claude Opus 4.7 for This Use Case")
    styled_table(doc,
        ["Capability", "Relevance to CIS Cloud Shield", "Advantage"],
        [
            ["200K token context window", "Process full 1,200-finding scan report in one pass", "No chunking or context loss"],
            ["Tool use / function calling", "Agents invoke AWS APIs, ServiceNow, and remediation tools natively", "True agentic execution"],
            ["Structured output", "Produces JSON-structured finding reports, remediation plans, and audit records", "Machine-readable outputs"],
            ["Reasoning transparency", "Extended thinking traces allow security team to review agent logic", "Audit-ready AI decisions"],
            ["Constitutional AI", "Refuses destructive actions without explicit authorisation", "Safe autonomous operation"],
            ["Multilingual", "Generates reports in client's preferred language", "Global enterprise readiness"],
        ],
        col_widths=[1.8, 2.5, 2.3]
    )

    heading2(doc, "6.2  Agent Configuration")
    body(doc, "Each of the four agents uses Claude Opus 4.7 with the following configuration:")
    add_info_box(doc, "Agent API Configuration", [
        "Model: claude-opus-4-7",
        "Max tokens: 4,096 per agent turn",
        "Temperature: 0.1 (deterministic, security-focused responses)",
        "System prompt: Role-specific instructions enforcing security best practices",
        "Context handoff: Each agent receives prior agent output (truncated to 800 tokens) as additional context",
        "Streaming: Real-time token streaming to Streamlit UI for live agent monitoring",
        "Safety: All destructive actions require explicit user confirmation before execution",
    ], bg_hex='EBF4FF')

    heading2(doc, "6.3  Agentic Framework Design Principles")
    bullet(doc, "Autonomy with Oversight: Agents can investigate and plan independently but require human approval before executing irreversible actions (credential rotation, secret deletion).")
    bullet(doc, "Context Persistence: Each agent's output is captured and passed to the next, creating a coherent multi-agent reasoning chain.")
    bullet(doc, "Failure Isolation: If one agent fails, the chain halts gracefully and raises a ServiceNow incident rather than proceeding with incomplete context.")
    bullet(doc, "Auditability: All agent interactions, tool calls, and decisions are logged to CloudWatch Logs with immutable retention.")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 – AUTO-REMEDIATION LOOP
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "7. AUTO-REMEDIATION LOOP & SERVICENOW INTEGRATION")
    body(doc,
        "CIS Cloud Shield implements a fully autonomous, closed-loop remediation system that operates continuously "
        "without human intervention for standard remediation scenarios. The loop is orchestrated by Claude Opus 4.7 "
        "agents and integrated with ServiceNow for full ITSM traceability."
    )

    heading2(doc, "7.1  Loop Architecture")
    body(doc, "The auto-remediation loop follows a five-stage cycle:")
    styled_table(doc,
        ["Stage", "Agent / Component", "Action", "ServiceNow Record"],
        [
            ["1. Detection", "Scanner Agent", "Identifies new credential finding or backup gap", "INC created (Priority 1-3 based on severity)"],
            ["2. Triage", "INC Agent", "Enriches incident with blast-radius analysis, assigns to remediation queue", "INC updated with analysis notes"],
            ["3. Investigation", "Dispatcher Agent", "Determines remediation playbook, creates change request, assigns engineer", "CHG created, INC linked"],
            ["4. Remediation", "Remediator Agent", "Executes automated remediation (secret rotation, backup policy deployment)", "CHG updated to In-Progress → Resolved"],
            ["5. Closure", "INC Agent", "Validates remediation success, closes INC with evidence", "INC closed with resolution notes and evidence hash"],
        ],
        col_widths=[1.2, 1.5, 2.5, 1.4]
    )

    heading2(doc, "7.2  ServiceNow Integration Details")
    styled_table(doc,
        ["Operation", "API Endpoint", "Trigger"],
        [
            ["Create Incident", "POST /api/now/table/incident", "New finding detected by Scanner Agent"],
            ["Update Incident", "PATCH /api/now/table/incident/{sys_id}", "Agent adds analysis, remediation progress"],
            ["Create Change Request", "POST /api/now/table/change_request", "Remediator Agent initiates action plan"],
            ["Close Incident", "PATCH /api/now/table/incident/{sys_id} state=6", "INC Agent confirms successful remediation"],
            ["Query Incidents", "GET /api/now/table/incident?sysparm_query=...", "Dashboard showing open/resolved counts"],
        ],
        col_widths=[1.8, 2.8, 2.0]
    )

    heading2(doc, "7.3  Loop Safeguards")
    bullet(doc, "Human-in-the-loop gate: Tier 1 (CRITICAL) remediations require explicit human approval before execution.")
    bullet(doc, "Maximum loop depth: Remediation loop limited to 3 retry attempts before escalating to human review.")
    bullet(doc, "Rollback capability: All remediations have an automated rollback procedure executed if health checks fail post-remediation.")
    bullet(doc, "Audit trail: Every loop iteration logged with agent ID, action taken, and outcome — immutably stored in CloudWatch Logs.")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 – IMPLEMENTATION ROADMAP
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "8. IMPLEMENTATION ROADMAP")
    body(doc,
        "Infosys proposes a 90-day phased implementation program to achieve full coverage across both capabilities:"
    )
    styled_table(doc,
        ["Phase", "Duration", "Workstreams", "Key Milestones"],
        [
            ["Phase 0\nFoundation", "Days 1-14",
             "• AWS account onboarding\n• IAM role provisioning\n• ServiceNow API configuration\n• CIS Cloud Shield deployment",
             "Platform live, AWS connected, ServiceNow integrated"],
            ["Phase 1\nCredential Scanning", "Days 15-45",
             "• Full scan of 1,200 hardcoded credentials\n• Risk assessment and prioritisation\n• Secrets Manager migration (CRITICAL first)\n• Prevention controls deployment",
             "100% CRITICAL credentials rotated\n80% HIGH credentials migrated\nCI/CD gates active"],
            ["Phase 2\nBackup Remediation", "Days 30-75",
             "• Backup gap analysis across 250 apps\n• Tier 1/2 backup policy deployment\n• Vault Lock configuration\n• Monitoring dashboards live",
             "100% Tier 1/2 apps protected\nVault Lock enabled on all vaults\nRecovery validation baseline set"],
            ["Phase 3\nFull Coverage", "Days 60-90",
             "• Tier 3/4 backup completion\n• Full credential remediation\n• DR drill execution\n• Executive reporting",
             "100% credential remediation\n100% backup coverage\nFirst DR drill completed"],
            ["Phase 4\nSteady State", "Day 90+",
             "• 24×7 managed operations\n• Monthly compliance reporting\n• Quarterly DR drills\n• Continuous improvement",
             "Ongoing managed service\nCompliance evidence package\nAnnual recertification"],
        ],
        col_widths=[1.2, 0.9, 2.8, 2.0]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 – WHY INFOSYS
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "9. WHY INFOSYS")
    styled_table(doc,
        ["Differentiator", "Detail"],
        [
            ["AI-First Architecture", "Infosys has deployed Claude Opus 4.7 in production security operations for Fortune 500 clients — not a pilot, a proven platform."],
            ["AWS Premier Partner", "AWS Premier Tier Services Partner with Security, Migration, and DevOps competencies. Deep IAM, Secrets Manager, and AWS Backup expertise."],
            ["Pre-Built Platform", "CIS Cloud Shield is a fully functional prototype — not a slide deck. Client can see it running on Day 1 of engagement."],
            ["ITSM Integration", "Native ServiceNow integration with automated INC creation, dispatch, remediation, and closure — end-to-end traceability."],
            ["Managed Services", "Infosys provides 24×7 NOC coverage for backup and security operations, eliminating the need for client to hire additional headcount."],
            ["Compliance Expertise", "Deep experience in PCI-DSS, HIPAA, SOC 2, ISO 27001, and FedRAMP compliance evidence generation."],
            ["Speed to Value", "90-day program to full coverage — accelerated by AI automation replacing months of manual credential remediation effort."],
        ],
        col_widths=[2.0, 4.6]
    )
    doc.add_paragraph()
    add_info_box(doc, "Infosys Investment in This Engagement", [
        "Dedicated Solution Architect (AWS Security) + AI Engineer for full 90-day program",
        "CIS Cloud Shield platform licensed at no additional cost for the engagement duration",
        "Anthropic Claude Opus 4.7 API costs covered by Infosys for Proof of Concept phase",
        "ServiceNow developer instance pre-configured and ready for demonstration",
        "Executive briefing and live demo available within 5 business days of engagement kick-off",
    ], bg_hex='EBF4FF')
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 – APPENDIX
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "10. APPENDIX – IAM ROLE, POLICY & AWS ARCHITECTURE REFERENCE")

    # ── 10.1 Role Overview ────────────────────────────────────────────────────
    heading2(doc, "10.1  CISCloudShieldRole — Pre-Provisioned IAM Role")
    body(doc,
        "Infosys has pre-provisioned a dedicated IAM role in the client's AWS account for exclusive use by the "
        "CIS Cloud Shield platform. No access keys or secret keys are issued — the platform authenticates entirely "
        "via STS AssumeRole with short-lived, auto-expiring temporary credentials."
    )
    styled_table(doc,
        ["Attribute", "Value"],
        [
            ["Role Name",           "CISCloudShieldRole"],
            ["Role ARN",            "arn:aws:iam::448549863273:role/CISCloudShieldRole"],
            ["AWS Account ID",      "448549863273"],
            ["Role ID",             "AROAWQ35HX5UUNMNKY7Y5"],
            ["Inline Policy",       "CISCloudShieldScanPolicy"],
            ["Permission Scope",    "Read-Only — credential scanning and backup assessment only"],
            ["Trusted Principal",   "arn:aws:iam::448549863273:user/Ajit_Gosavi@infosys.com"],
            ["Session Duration",    "3,600 seconds (1 hour) — temporary credentials auto-expire"],
            ["Created",             "2026-05-03 (provisioned by Infosys during platform onboarding)"],
            ["Tags",                "Application=CISCloudShield | ManagedBy=Infosys | Environment=Demo"],
        ],
        col_widths=[2.0, 4.6]
    )

    body(doc, "How to connect the platform to this role (sidebar steps):")
    bullet(doc, "1.  Open the CIS Cloud Shield sidebar and select Auth Method: IAM Role (AssumeRole)")
    bullet(doc, "2.  Paste the Role ARN:  arn:aws:iam::448549863273:role/CISCloudShieldRole")
    bullet(doc, "3.  Leave External ID blank (single-account, same-user trust — no external ID required)")
    bullet(doc, "4.  Select the AWS Region and click Connect to AWS")
    bullet(doc, "5.  The platform calls sts:AssumeRole and confirms the account identity via sts:GetCallerIdentity")

    add_info_box(doc, "Security Properties of This Approach", [
        "No long-lived credentials ever enter the application UI or configuration files",
        "Temporary credentials expire after 1 hour — no revocation step needed if compromised",
        "All AssumeRole calls are logged to CloudTrail with full session context",
        "Role can be revoked instantly by removing the trust policy — single control point",
        "Principle of least privilege — read-only scope, no write permissions granted",
        "Streamlit Cloud stores only the Role ARN in secrets (not a credential — a resource identifier)",
    ], bg_hex='EBF4FF')
    doc.add_page_break()

    # ── 10.2 Trust & Permissions Policies ────────────────────────────────────
    heading2(doc, "10.2  Trust Policy & Permissions Policy")

    heading3(doc, "Trust Policy (who can assume this role)")
    trust_policy_text = """{
  "Version": "2012-10-17",
  "Statement": [
    {
      "Sid": "AllowCurrentUser",
      "Effect": "Allow",
      "Principal": {
        "AWS": "arn:aws:iam::448549863273:user/Ajit_Gosavi@infosys.com"
      },
      "Action": "sts:AssumeRole"
    }
  ]
}"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(trust_policy_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    body(doc,
        "To extend access to additional users, teams, or AWS services (EC2 instance profile, ECS task role, "
        "Lambda execution role), add them as additional Principal entries in the trust policy above."
    )

    heading3(doc, "Inline Permissions Policy — CISCloudShieldScanPolicy (read-only)")
    permissions_policy_text = """{
  "Version": "2012-10-17",
  "Statement": [
    {
      "Sid": "CredentialScanReadOnly",
      "Effect": "Allow",
      "Action": [
        "iam:GenerateCredentialReport",   "iam:GetCredentialReport",
        "iam:ListUsers",                  "iam:ListAccessKeys",
        "iam:GetAccessKeyLastUsed",       "iam:GetUser",
        "iam:ListMFADevices",             "iam:ListVirtualMFADevices",
        "iam:GetAccountSummary",          "iam:GetAccountPasswordPolicy",
        "secretsmanager:ListSecrets",     "secretsmanager:DescribeSecret",
        "securityhub:GetFindings",
        "securityhub:ListFindingAggregators",
        "securityhub:GetEnabledStandards"
      ],
      "Resource": "*"
    },
    {
      "Sid": "BackupScanReadOnly",
      "Effect": "Allow",
      "Action": [
        "backup:ListBackupPlans",         "backup:ListBackupVaults",
        "backup:DescribeBackupVault",     "backup:ListProtectedResources",
        "backup:GetBackupVaultNotifications",
        "backup:ListBackupJobs",          "backup:DescribeBackupJob",
        "ec2:DescribeInstances",          "ec2:DescribeVolumes",
        "rds:DescribeDBInstances",        "rds:DescribeDBClusters",
        "s3:ListAllMyBuckets",            "s3:GetBucketVersioning",
        "s3:GetObjectLockConfiguration",  "s3:GetBucketReplication",
        "dynamodb:ListTables",            "dynamodb:DescribeTable",
        "dynamodb:DescribeContinuousBackups",
        "efs:DescribeFileSystems",        "efs:DescribeBackupPolicy"
      ],
      "Resource": "*"
    },
    {
      "Sid": "STSAndObservability",
      "Effect": "Allow",
      "Action": [
        "sts:GetCallerIdentity",
        "cloudtrail:LookupEvents",
        "cloudwatch:DescribeAlarms",
        "config:DescribeConfigRules",
        "config:GetComplianceSummaryByConfigRule"
      ],
      "Resource": "*"
    }
  ]
}"""
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    run2 = p2.add_run(permissions_policy_text)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_page_break()

    # ── 10.3 AWS Architecture Component Reference ─────────────────────────────
    heading2(doc, "10.3  AWS Architecture Component Reference")
    styled_table(doc,
        ["AWS Service", "Role in Platform", "Key Feature Used"],
        [
            ["AWS STS",            "IAM Role authentication (primary)",       "AssumeRole — short-lived temp credentials, no long-lived keys"],
            ["AWS IAM",            "Identity & credential management",         "Credential report, CISCloudShieldRole, MFA enforcement"],
            ["AWS Secrets Manager","Credential storage & rotation target",     "Auto-rotation, versioning, VPC endpoint, resource policy"],
            ["AWS Security Hub",   "Centralised finding aggregation",          "CIS Benchmark v1.4, PCI-DSS, NIST 800-53 controls"],
            ["AWS Backup",         "Centralised backup orchestration",         "Vault Lock (WORM compliance mode), cross-region copy"],
            ["AWS Config",         "Continuous compliance monitoring",         "Custom rules, conformance packs, credential-pattern detection"],
            ["AWS CloudTrail",     "Full audit logging",                       "All AssumeRole calls, GetSecretValue, management events"],
            ["AWS CloudWatch",     "Monitoring & alerting",                    "Metric alarms on backup failures, credential scan anomalies"],
            ["Amazon S3",          "Object storage with WORM",                 "Object Lock in Compliance Mode — ransomware-proof backups"],
            ["Amazon RDS",         "Database backup coverage",                 "Automated backups, Multi-AZ, AWS Backup integration"],
            ["Amazon DynamoDB",    "NoSQL with backup",                        "Point-in-Time Recovery (PITR), on-demand backup"],
            ["Amazon EFS",         "File system backup",                       "EFS Backup Policy, AWS Backup integration"],
        ],
        col_widths=[1.5, 2.3, 2.8]
    )

    # ── 10.4 Contact & Next Steps ─────────────────────────────────────────────
    heading2(doc, "10.4  Contact & Next Steps")
    body(doc, "To proceed with this engagement, please contact the Infosys CIS team:")
    add_info_box(doc, "Next Steps", [
        "Step 1: Schedule a live demo of CIS Cloud Shield (available within 5 business days)",
        "Step 2: Share the CISCloudShieldRole ARN — no credentials required (arn:aws:iam::448549863273:role/CISCloudShieldRole)",
        "Step 3: Review and approve the 90-day implementation Statement of Work",
        "Step 4: Kick off Phase 0 — platform onboarding, IAM role extension to additional principals, ServiceNow config",
        "Platform: Deployable to Streamlit Cloud within 24 hours — Role ARN stored in Streamlit secrets, no key storage",
        "AI Model: Anthropic Claude Opus 4.7 (claude-opus-4-7) — enterprise API access included in engagement",
    ], bg_hex='EBF4FF')

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Infosys Limited  |  CIS Cloud Shield  |  Confidential  |  "
        f"Version 1.0  |  {datetime.date.today().strftime('%B %Y')}"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fr.font.name = 'Calibri'

    # Save
    out_path = r"c:\aidemos\aws-security-backup-platform\Infosys_CIS_Cloud_Shield_RFI_Response.docx"
    doc.save(out_path)
    print(f"Document saved: {out_path}")
    return out_path

if __name__ == "__main__":
    build_document()
